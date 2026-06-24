#!/usr/bin/env python3
"""
DOCX文档解析模块 - Phase 6.1.2
"""

from typing import Dict, Optional

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

TABLE_ROW_MIN_CELLS = 2


def parse_docx(file_path: str) -> Dict:
    """完整的DOCX文本提取"""
    
    if not DOCX_AVAILABLE:
        raise Exception("❌ 缺少python-docx库。安装：pip install python-docx")
    
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        full_text = ""
        all_tables = []
        metadata = {"paragraphs": 0, "tables": 0, "sections": len(doc.sections)}
        
        # 提取页眉
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                if para.text.strip():
                    full_text += f"[页眉] {para.text}\n"
        
        # 提取主文本和表格
        for element in doc.element.body:
            if element.tag.endswith('p'):
                para_obj = None
                for para in doc.paragraphs:
                    if para._element == element:
                        para_obj = para
                        break
                
                if para_obj:
                    text = para_obj.text.strip()
                    if text:
                        full_text += text + "\n"
                        metadata["paragraphs"] += 1
            
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        table_data = _extract_docx_table(table)
                        if table_data:
                            all_tables.append(table_data)
                            metadata["tables"] += 1
                        break
        
        # 提取页脚
        for section in doc.sections:
            footer = section.footer
            for para in footer.paragraphs:
                if para.text.strip():
                    full_text += f"[页脚] {para.text}\n"
        
        core_props = doc.core_properties
        
        return {
            "text": full_text.strip(),
            "tables": all_tables,
            "metadata": {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "paragraphs": metadata["paragraphs"],
                "tables": metadata["tables"],
                "sections": metadata["sections"]
            },
            "quality": {
                "has_tables": len(all_tables) > 0,
                "confidence": 0.95,
                "method": "python-docx"
            }
        }
    
    except Exception as e:
        raise Exception(f"DOCX提取失败：{str(e)}")


def _extract_docx_table(table) -> Optional[Dict]:
    """从DOCX表格提取数据"""
    
    try:
        table_data = {"rows": []}
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            
            if any(cell for cell in row_data if cell):
                table_data["rows"].append(row_data)
        
        if len(table_data["rows"]) >= 2 and any(len(row) >= TABLE_ROW_MIN_CELLS for row in table_data["rows"]):
            return table_data
        
        return None
    
    except Exception:
        return None

